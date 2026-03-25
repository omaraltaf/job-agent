from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def insert_hr(paragraph):
    """Inserts a horizontal rule (bottom border) to the given paragraph."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'A0A0A0')  # Light gray
    pBdr.append(bottom)
    pPr.append(pBdr)

def markdown_to_docx(markdown_text: str, output_path: str):
    doc = Document()
    
    # Basic styling setup to make it look professional
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    
    # Adjust margins to fit more content (standard CV trick)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    lines = markdown_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        # Intercept formatting errors gracefully
        
        # 1. Catch horizontal dividers (e.g. ---, --, ***)
        if re.match(r'^[-_*]{2,}\s*$', line):
            p = doc.add_paragraph()
            insert_hr(p)
            continue
            
        # 2. Catch Headings
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            insert_hr(p)  # Professional CVs usually have a line under Section Headers
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            
        # 3. Catch Normal Bullets
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            text_content = line[2:].strip()
            _add_formatted_text(p, text_content)
            
        # 4. Catch Claude's buggy '--' or bad dashes
        elif line.startswith('-') or line.startswith('--'):
            clean_line = re.sub(r'^[-]+', '', line).strip()
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, clean_line)
            
        # 5. Normal paragraphs
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, line)
            
    doc.save(str(output_path))

def _add_formatted_text(paragraph, text):
    """Basic handling for **bold** and *italic* text from markdown."""
    # First, handle bold
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Then handle italics within the non-bold parts
            subparts = re.split(r'(\*.*?\*)', part)
            for subpart in subparts:
                if subpart.startswith('*') and subpart.endswith('*') and not subpart.startswith('**'):
                    run = paragraph.add_run(subpart[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(subpart)
